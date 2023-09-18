from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import FormattedText
from pdf2image import convert_from_path
import time
from docx.shared import Cm, Pt, Inches
from docx import Document
from docx.enum.section import WD_SECTION_START


class PageInFile:
    # 아래 코드로 바꿔야 함
    # def __init__(self, image_list, column_num):
    #     self.image_list = image_list
    def __init__(self, pdf_url, column_num):
        self.pdf_url = pdf_url
        self.image_list = self.pdf2jpg(pdf_url)
        self.column_num = column_num

        # 이거 create_text_list call한 값으로 check_connecting_line 호출하면 될듯
        self.image_extracted_data_list = self.create_text_list(self.image_list, self.column_num)

    def pdf2jpg(self, pdf_url):
        images = list()
        if not pdf_url.endswith('.pdf'):
            print("This is not a PDF file")
        else:
            images = convert_from_path(pdf_url)
        return images

    def create_text_list(self, images, column_num):
        image_list = list()

        for idx, image in enumerate(images):
            # image_numpy = np.array(image)
            text_in_image = FormattedText.FormattedText(image, column_num[idx])
            image_list.append(text_in_image)


        # 여기서 list를 check_connecting_line에 넣어주면 될듯
        # 받아오는 값을 image_extracted_data_list에 넣어주면 될듯


        return image_list

    def points_to_cm(self, points, height):
        return (points / height) * 29.7

        # 불러올 때, self.insert_text_to_word(image_list, self.pdf_url[:-4]+".docx")
    def insert_text_to_word(self, image_list, output_filename):
        # Create a new Word document
        doc = Document()
        sections = doc.sections
        for idx, page in enumerate(image_list):
            # 일단은 다음줄 연결 없이 한 장씩 뽑기
            # new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            # new_section.start_type
            new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            new_section.start_type

            new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            new_section.start_type


            for idx, par in enumerate(page.line_formatted_data):
                new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
                new_section.start_type
                if isinstance(par, dict):
                    left = Pt(par['min_x'])
                    top = Pt(par['min_y'])
                    width = Pt(par['width'])
                    height = Pt(par['height'])


                    # shape = doc.add_shape(1, left, top, width, height)
                    # text_frame = textbox.text_frame
                    # shape.fill.solid()
                    # shape.fill.fore_color.rgb = (255,0,0)
                    # shape.line.width = Pt(2)
                    # shape.line.color.rgb = (0,0,0)



                    if par['text'].startswith('e '):
                        p = doc.add_run(par['text'][2:-2], style='List Bullet')
                    else:
                        p = doc.add_run(par['text'][:-2])
                    p.font.size = Pt(12)

                    # spacing_after = Pt(0)  # Adjust this value as needed
                    # p.space_after = spacing_after

                    # 아래 alignment 코드는 google docs에서만 작동함.
                    # p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    # style = doc.styles['Normal']
                    # font = style.font
                    # font.size = Pt(10)

        doc.save(output_filename)


if __name__ == '__main__':
    start = time.time()

    extractor = PageInFile("./test/rt_img.pdf", [2])
    # extractor = PageInFile("C:/Users/SOL/Documents/sandBox/JpgToDocs/test/rt_img.pdf", [2])

    extractor.insert_text_to_word(extractor.image_extracted_data_list, extractor.pdf_url[:-4] + ".docx")
    end = time.time()

    print("\nIt took " + str(end - start) + " seconds")

