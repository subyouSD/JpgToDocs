from ImageTextExtractor import *
import docx
from docx.shared import Pt


def to_docx():
    doc = docx.Document()
    doc.add_heading('from pdf to docx')

    doc.add_heading('increased font size paragraph: ', 1)
    para = doc.add_paragraph().add_run(
        "ddddd"
    )

    para.font.size = Pt(12)
    doc.add_heading('normal font size paragraph: ', 1)
    doc.add_paragraph().add_run(
        "ddddd"
    )

    doc.save('to_docx1.docx')


def test():
    start = time.time()
    images = pdf2jpg("./test/Meeting Minute3.pdf")
    for i, image in enumerate(images):
        extracted_data = extract_data_from_image(image)

        line_formatted = (format_extracted_data_to_text(extracted_data))
        for j in line_formatted:
            print(j)
    end = time.time()

    print("\nit took " + str(end - start) + " seconds")


test()