from PIL import Image

import LineFormattedData
from pdf2image import convert_from_path

from docx.shared import Cm, Pt, Inches, Mm
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING


class TextInImages:
    # 아래 코드로 바꿔야 함
    # def __init__(self, image_list, column_num):
    #     self.image_list = image_list
    def __init__(self, pdf_url, column_num):
        self.pdf_url = pdf_url
        # self.image_list = [Image.open(pdf_url)]
        self.image_list = self.pdf2jpg(pdf_url)
        self.column_num = column_num
        self.section_height = 29.7

        # 이거 create_text_list call한 값으로 check_connecting_line 호출하면 될듯
        self.image_extracted_data_list = self.create_text_list(self.image_list, self.column_num)

    def pdf2jpg(self, pdf_url):
        images = list()
        if not pdf_url.endswith('.pdf'):
            print("This is not a PDF file")
        else:
            images = convert_from_path(pdf_url)
        return images

    def create_text_list(self, images, column_num):
        image_list = []

        for idx, image in enumerate(images):
            # image_numpy = np.array(image)
            text_in_image = LineFormattedData.LineFormattedData(image, column_num[idx])
            image_list.append(text_in_image)

        return image_list

    def points_to_cm(self, length, height):
        return length * self.section_height / height

    def insert_text_to_word(self, image_list, output_filename):
        # Create a new Word document
        doc = Document()
        sections = doc.sections
        section = sections[0]

        for idx, page in enumerate(image_list):
            print(page.height, page.width)
            if round((page.height / page.width), 1) == 1.4:
                section.page_width = Mm(210)
                section.page_height = Mm(297)
                self.section_height = 29.7
            if round((page.height / page.width), 1) == 1.3:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                self.section_height = 27.94

            if page.column_num == 1:
                for section in sections:
                    section.top_margin = Cm(self.points_to_cm(page.top_y, page.height))
                    section.bottom_margin = Cm(self.points_to_cm(page.height - page.bottom_y, page.height))
                    section.left_margin = Cm(self.points_to_cm(page.left_x, page.height))
                    section.right_margin = Cm(self.points_to_cm((page.width - page.right_x), page.height))

            elif page.column_num == 2:

                sectPr = section._sectPr
                cols = sectPr.xpath('./w:cols')[0]
                cols.set(qn('w:num'), '2')
                spacing_between_columns = self.points_to_cm((page.second_column_left_x - page.first_column_right_x),
                                           page.height) * 567

                cols.set(qn('w:space'), str(spacing_between_columns))

                section.top_margin = Cm(self.points_to_cm(page.top_y, page.height))
                section.bottom_margin = Cm(self.points_to_cm((page.height - page.bottom_y), page.height))
                section.left_margin = Cm(self.points_to_cm(page.first_column_left_x, page.height))
                section.right_margin = Cm(self.points_to_cm((page.width - page.second_column_right_x), page.height))

            for par in page.line_formatted_data:
                if par['text'].startswith('e '):
                    doc.add_paragraph(par['text'][2:], style='List Bullet')
                else:
                    doc.add_paragraph(par['text'])
                style = doc.styles['Normal']
                font = style.font
                font.size = Pt(9)
                style.paragraph_format.space_after = Pt(2)
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                style.paragraph_format.line_spacing = None

                style = doc.styles['List Bullet']
                style.paragraph_format.left_indent = Inches(0.15)
                style.paragraph_format.first_line_indent = Inches(-0.15)


        doc.save(output_filename)



