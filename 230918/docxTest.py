
from docx.shared import Cm, Pt, Inches, Mm
from docx import Document
from docx.oxml.ns import qn


line_formatted_data = [{'text': '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n', 'x': 0, 'y': 0},
                       {'text': 'Fig.3 Scientific Articles from year 2009 to 2021 ', 'x': 143, 'y': 795},
                       {'text': 'proceedings, and databases were investigated and explored in the area of homomorphic encryption. Scientific publi cations from Scopus, ACM Digital library, Springer Link, Sciencedirect, Google Scholar were selected based upon research questions given in Table 1. Fundamental studies in PHE, SWHE that function as cornerstones of homomor phic encryption were chosen first, followed by publications focused only on fully homomorphic encryption. A total of 1815 scientific papers were obtained using a database search with Keywords used for search as *>homomorphic encryp tion’, and *homomorphic encryption’ or “medical” OR “healthcare” OR “bioinformatics” OR “EHR” OR “patient” OR “health” OR “medicine”. Following the removal of dupli cate documents, a total of 857 records were evaluated for title screening. After screening title total 194 articles were selected for abstract screening. In abstract screening 69 papers that were focused on homomorphic encryption based upon LWE, NTRU, Lattices, Integers, and HE papers focused mainly on healthcare and bioinformatics were selected. Other than that 19 important papers in PHE and SWHE (consider Fig. 4) were chosen for better understanding the concept of homomorphic encryption. \n', 'x': 142, 'y': 853},
                       {'text': 'Inclusion and exclusion criteria \n', 'x': 144, 'y': 1651},
                       {'text': 'Scientific articles in journals, conferences proceedings, workshops published in the year range from Jan 2009 to Dec 2021 (Fig.3) were considered. Other than that, Partial and Somewhat homomorphic encryption papers of the old era were also included to better understand the concept of homo morphic encryption. All quality research publications of fully homomorphic encryption after the Gentry FHE scheme were considered. Articles that were focused on the applicability of homomorphic encryption other than healthcare or bioinfor matics were excluded. ', 'x': 143, 'y': 1723},
                       {'text': 'Research questions \n', 'x': 852, 'y': 232},
                       {'text': 'Prime objective of this review writing is to categorize the current literature on homomorphic encryption with its con tributions in health informatics. This research study’s end result is the identification and examination of homomorphic encryption methods in healthcare. A set of research questions are formulated for this systematic literature review in Table ', 'x': 851, 'y': 304},
                       {'text': '1 \n', 'x': 854, 'y': 513},
                       {'text': 'Background \n', 'x': 852, 'y': 611},
                       {'text': 'The medical services industry is experiencing a digital rev olution. Modernizing medical care has prompted another time of computerized wellbeing and health. Medical ser vices information is gathered from different sources (e.g., sensors associated with patients) and stored in unique med ical services clouds (e.g., private and public clouds). Also, the volume of agglomerated medical information is suffi ciently enormous to qualify as “Big Data”. As cloud medical services become a well-defined component in the medi cal services industry, there is a more critical requirement for safely sharing patient data across such dissimilar med ical services clouds. Besides, with Accountable Treatment Organizations (ACOs) (e.g., medical service providers, spe cialists, clinics, and protection providers) collaborate to provide top-notch care, with demand for constant availabil ity across cloud medical services higher than at any point in recent times. A disentangled patient-driven paradigm, in which patients can switch suppliers while still providing their data in a useful way for better diagnosis and treatment, and, in the long term, for enhanced global health, is appeal ing. As of now, medical care suppliers who have delicate patient information in private medical care clouds across the globe are reluctant to share that data on account of secu rity and protection issues. As medical care suppliers move to the local area and public cloud-based administrations, a requirement for a secure connection between divergent med ical care cloud increments. Moreover, security guidelines forced by Health Insurance Portability and Accountability Act (HIPAA) [6] and Health Information Technology for Economic and Clinical Health (HITECH) [7] place a cumber some undertaking on medical care Information Technology (IT) framework to be agreeable with protection and security guidelines. Moreover, with arising Internet of Things (IoT) market and its mix in the vast information cloud stage, there is expanded worry about security and protection with the med ical services cloud worldview. Many researchers contributed with their study in homomorphic encryption. Homomorphic encryption has three types: partial homomorphic encryption (PHE), somewhat homomorphic encryption (SWHE), and fully homomorphic encryption Fig.5. PHE supports either', 'x': 851, 'y': 685}]


def insert_text_to_word(line_formatted_data, output_filename):
    # 페이지 사이즈 비율에 따라, A4용지 또는 US Letter 용지를 선택




    doc = Document()
    sections = doc.sections

    section = sections[0]
    if round(2339 / 1653, 1) == 1.4:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    if round(2339 / 1653, 1) == 1.3:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    spacing_between_columns = (54/2339)*29.7*567
    cols.set(qn('w:space'),str(spacing_between_columns))

    section.top_margin = Cm((232*29.7/2339)-1.27)
    section.bottom_margin = Cm(((2339-2074.12)*29.7/2339)-1.27)

    section.left_margin = Cm(142 * 21/1653)
    section.right_margin = Cm((1653-1512)*21/1653)

    for par in line_formatted_data:
        if isinstance(par, dict):
            if par['text'].startswith('e '):
                doc.add_paragraph(par['text'][2:], style='List Bullet')
            else:
                doc.add_paragraph(par['text'])
            style = doc.styles['Normal']
            font = style.font
            font.size = Pt(9.5)

    doc.save(output_filename)

insert_text_to_word(line_formatted_data, "test.docx")
