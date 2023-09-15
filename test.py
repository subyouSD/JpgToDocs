from docx import Document

from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn


doc = Document()
# First Section
new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
new_section.start_type

p1 = doc.add_heading('This is the title!!!!!!!!!!!!!')
p1.alignment = 1

# Second Section
new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
new_section.start_type
section = doc.sections[2]

# Set to 2 column layout
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')

p1 = doc.add_paragraph('This is the 1st para')
p1.alignment = 3

# Third Section
new_section = doc.add_section(WD_SECTION_START.NEW_COLUMN)
new_section.start_type


p1 = doc.add_paragraph('This is the 2nd para')
p1.alignment = 3


# Save
doc.save('demo.docx')