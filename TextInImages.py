import LineFormattedData
from pdf2image import convert_from_path
import time
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

class TextInImages:
    # 아래 코드로 바꿔야 함
    # def __init__(self, image_list, column_num):
    #     self.image_list = image_list
    def __init__(self, pdf_url, column_num):
        self.pdf_url = pdf_url
        self.image_list = self.pdf2jpg(pdf_url)
        self.column_num = column_num

        # 이거 create_text_list call한 값으로 check_connecting_line 호출하면 될듯
        self.image_extracted_data_list = self.create_text_list(self.image_list, self.column_num)

    def pdf2jpg(self, pdf_url):
        images = list()
        if not pdf_url.endswith('.pdf'):
            print("This is not a PDF file")
        else:
            images = convert_from_path(pdf_url)
        return images

    def create_text_list(self, images, column_num):
        image_list = []

        for idx, image in enumerate(images):
            # image_numpy = np.array(image)
            text_in_image = LineFormattedData.LineFormattedData(image, column_num[idx])
            image_list.append(text_in_image)


        # 여기서 list를 check_connecting_line에 넣어주면 될듯
        # 받아오는 값을 image_extracted_data_list에 넣어주면 될듯


        return image_list

    def points_to_cm(self, points, height):
        return Cm(points / height * 29.7)

        # 불러올 때, self.insert_text_to_word(image_list, self.pdf_url[:-4]+".docx")
    def insert_text_to_word(self, image_list, output_filename):
        # Create a new Word document
        doc = Document()
        sections = doc.sections
        for idx, page in enumerate(image_list):
            # 일단은 다음줄 연결 없이 한 장씩 뽑기
            # Create a new paragraph
            new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            new_section.start_type


            # p1 = doc.add_heading('This is the title!!!!!!!!!!!!!')
            # p1.alignment = 1

            # Second Section
            new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            new_section.start_type

            section = sections[2]

            if page.column_num == 1:
                for section in sections:
                    section.top_margin = self.points_to_cm(page.top_y, page.height)
                    section.bottom_margin = self.points_to_cm(page.height - page.bottom_y, page.height)
                    section.left_margin = self.points_to_cm(page.left_x, page.height)
                    section.right_margin = self.points_to_cm(page.left_x, page.height)

            elif page.column_num == 2:
            # # Set to 2 column layout
                sectPr = section._sectPr
                cols = sectPr.xpath('./w:cols')[0]
                cols.set(qn('w:num'), '2')

            # 1440 twip units are 1 inch.
                space_value = 0.5 * 1440
                cols.set(qn('w:space'), str(space_value))

                for section in sections:
                # for section in sections:
                    section.top_margin = self.points_to_cm(page.top_y, page.height)
                    section.bottom_margin = self.points_to_cm(page.height - page.bottom_y, page.height)
                    section.left_margin = self.points_to_cm(page.first_column_left_x, page.height)
                    section.right_margin = self.points_to_cm(page.first_column_left_x, page.height)


            # Third Section
            # new_section = doc.add_section(WD_SECTION_START.NEW_COLUMN)
            # new_section.start_type

            # p1 = doc.add_paragraph('This is the 2nd para')
            # p1.alignment = 1

            # paragraph = doc.add_paragraph()
            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            for par in page.line_formatted_data:
                if isinstance(par, dict):
                    if par['text'].startswith('e '):
                        p = doc.add_paragraph(par['text'][2:], style='List Bullet')
                        # p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    # Add text with the specified font size
                    else:
                        # run = paragraph.add_run(par['text'])  # Use .get() to handle missing 'text' key
                        p = doc.add_paragraph(par['text'])
                        # p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    style = doc.styles['Normal']
                    font = style.font
                    font.size = Pt(10)

        doc.save(output_filename)


if __name__ == '__main__':
    start = time.time()

    # extractor = TextInImages("./test/rt_img.pdf", [2])
    extractor = TextInImages("C:/Users/SOL/Documents/sandBox/JpgToDocs/test/rt_img.pdf", [2])

    extractor.insert_text_to_word(extractor.image_extracted_data_list, extractor.pdf_url[:-4] + ".docx")
    # print(extractor.pdf_url)
    # for i in extractor.image_extracted_data_list:
    #     for j in i:
    #         # print(j['text'], end="")
    #         print(j)
    end = time.time()

    print("\nIt took " + str(end - start) + " seconds")

